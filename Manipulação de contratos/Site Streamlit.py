import streamlit as st
import os
from docx import Document as doc
con_clube = doc('Manipulação de contratos\Contratos\CONTRATO - CLUBE.docx')
con_domestica = doc('Manipulação de contratos/Contratos/CONTRATO - DOMESTICA.docx')
con_padrao = doc('Manipulação de contratos/Contratos/CONTRATO - PADRAO.docx')
con_permuta = doc('Manipulação de contratos/Contratos/CONTRATO - PERMUTA.docx')

st.set_page_config(page_title='Preencher Contratos Fácil e Rápido.com.br')
listaaquivos = os.listdir('Manipulação de contratos\Contratos')

st.subheader('Contratos disponíveis:')
ce = st.selectbox(label='Contratos:', options=listaaquivos)

RAZAO_SOCIAL = st.text_input('RAZAO_SOCIAL')
CNPJ_CONTRATANTE = st.text_input('CNPJ_CONTRATANTE')
CEP_EMPRESA = st.text_input('CEP_EMPRESA')
ENDERECO_EMPRESA = st.text_input('ENDERECO_EMPRESA')
NOME_AVALISTA = st.text_input('NOME_AVALISTA')
CPF_AVALISTA = st.text_input('CPF_AVALISTA')
ENDERECO_AVALISTA = st.text_input('ENDERECO_AVALISTA')
CEP_AVALISTA = st.text_input('CEP_AVALISTA')
VALOR_MENSAL = st.text_input('VALOR_MENSAL')
VALOR_POR_EXTENSO = st.text_input('VALOR_POR_EXTENSO')
QTD_FUNC = st.text_input('QTD_FUNC')
VALOR_FUNC = st.text_input('VALOR_FUNC')
VALOR_POR_EXTENSO_FUNC = st.text_input('VALOR_POR_EXTENSO_FUNC')
REGIME = st.text_input('REGIME')
QTD_FISCAL = st.text_input('QTD_FISCAL')
QTD_CONTABIL = st.text_input('QTD_CONTABIL')
VENCIMENTO = st.text_input('VENCIMENTO')
DATA_VIGENCIA = st.text_input('DATA_VIGENCIA')
DATA_ASSINATURA = st.text_input('DATA_ASSINATURA')

campos_contrato = [
RAZAO_SOCIAL,
CNPJ_CONTRATANTE,
CEP_EMPRESA,
ENDERECO_EMPRESA,
NOME_AVALISTA,
CPF_AVALISTA,
ENDERECO_AVALISTA,
CEP_AVALISTA,
VALOR_MENSAL,
VALOR_POR_EXTENSO,
QTD_FUNC,
VALOR_FUNC,
VALOR_POR_EXTENSO_FUNC,
REGIME,
QTD_FISCAL,
QTD_CONTABIL,
VENCIMENTO,
DATA_VIGENCIA,
DATA_ASSINATURA
]
campos_contrato_strings =[
'RAZAO_SOCIAL',
'CNPJ_CONTRATANTE',
'CEP_EMPRESA',
'ENDERECO_EMPRESA',
'NOME_AVALISTA',
'CPF_AVALISTA',
'ENDERECO_AVALISTA',
'CEP_AVALISTA',
'VALOR_MENSAL',
'VALOR_POR_EXTENSO',
'QTD_FUNC',
'VALOR_FUNC',
'VALOR_POR_EXTENSO_FUNC',
'REGIME',
'QTD_FISCAL',
'QTD_CONTABIL',
'VENCIMENTO',
'DATA_VIGENCIA',
'DATA_ASSINATURA'
]
campos_contrato_final ={
'RAZAO_SOCIAL': '',
'CNPJ_CONTRATANTE' : '',
'CEP_EMPRESA': '',
'ENDERECO_EMPRESA' : '', 
'NOME_AVALISTA' : '',
'CPF_AVALISTA' : '',
'ENDERECO_AVALISTA' : '',
'CEP_AVALISTA' : '',
'VALOR_MENSAL' : '',
'VALOR_POR_EXTENSO' : '',
'QTD_FUNC' : '',
'VALOR_FUNC' : '',
'VALOR_POR_EXTENSO_FUNC' : '',
'REGIME' : '',
'QTD_FISCAL' : '',
'QTD_CONTABIL' : '',
'VENCIMENTO' : '',
'DATA_VIGENCIA' : '',
'DATA_ASSINATURA' : ''
}
col1, col2 = st.columns([0.4,1])

def contrato_clube():
    with col1:
        y = 0
        preencher = st.button(label='Preencher Campos')
        if preencher == True:
            for campo in campos_contrato:
                if campo == '':
                    st.error('Todos os campos devem estar preenchidos.')
                else:
                    campos_contrato_final[campos_contrato_strings[y]] = (campo)
                    y += 1
            x = 0
            for paragrafo in con_clube.paragraphs:
                if '{{' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace("{{", '')
                if '}}' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace("}}", '')
                if 'RAZAO_SOCIAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('RAZAO_SOCIAL', f"{campos_contrato_final['RAZAO_SOCIAL']}")
                if 'CNPJ_CONTRATANTE' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CNPJ_CONTRATANTE', f"{campos_contrato_final['CNPJ_CONTRATANTE']}")
                if 'CEP_EMPRESA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CEP_EMPRESA', f"{campos_contrato_final['CEP_EMPRESA']}")
                if 'ENDERECO_EMPRESA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('ENDERECO_EMPRESA', f"{campos_contrato_final['ENDERECO_EMPRESA']}")
                if 'NOME_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('NOME_AVALISTA', f"{campos_contrato_final['NOME_AVALISTA']}")
                if 'CPF_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CPF_AVALISTA', f"{campos_contrato_final['CPF_AVALISTA']}")
                if 'ENDERECO_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('ENDERECO_AVALISTA', f"{campos_contrato_final['ENDERECO_AVALISTA']}")
                if 'CEP_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CEP_AVALISTA', f"{campos_contrato_final['CEP_AVALISTA']}")
                if 'VALOR_MENSAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_MENSAL', f"{campos_contrato_final['VALOR_MENSAL']}")
                if 'VALOR_POR_EXTENSO' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_POR_EXTENSO', f"{campos_contrato_final['VALOR_POR_EXTENSO']}")
                if 'QTD_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_FUNC', f"{campos_contrato_final['QTD_FUNC']}")
                if 'VALOR_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_FUNC', f"{campos_contrato_final['VALOR_FUNC']}")
                if 'VALOR_POR_EXTENSO_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_POR_EXTENSO_FUNC', f"{campos_contrato_final['VALOR_POR_EXTENSO_FUNC']}")
                if 'REGIME' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('REGIME', f"{campos_contrato_final['REGIME']}")
                if 'QTD_FISCAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_FISCAL', f"{campos_contrato_final['QTD_FISCAL']}")
                if 'QTD_CONTABIL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_CONTABIL', f"{campos_contrato_final['QTD_CONTABIL']}")
                if 'VENCIMENTO' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VENCIMENTO', f"{campos_contrato_final['VENCIMENTO']}")
                if 'DATA_VIGENCIA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('DATA_VIGENCIA', f"{campos_contrato_final['DATA_VIGENCIA']}")
                if 'DATA_ASSINATURA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('DATA_ASSINATURA', f"{campos_contrato_final['DATA_ASSINATURA']}")
                if x < 18:
                    x += 1

            con_clube.save("Manipulação de contratos\ CONTRATO  PREENCHIDO- CLUBE.docx")
def contrato_domestica():
    with col1:
        y = 0
        preencher = st.button(label='Preencher Campos')
        if preencher == True:
            for campo in campos_contrato:
                if campo == '':
                    st.error('Todos os campos devem estar preenchidos.')
                else:
                    campos_contrato_final[campos_contrato_strings[y]] = (campo)
                    y += 1
            x = 0
            for paragrafo in con_domestica.paragraphs:
                if '{{' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace("{{", '')
                if '}}' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace("}}", '')
                if 'RAZAO_SOCIAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('RAZAO_SOCIAL', f"{campos_contrato_final['RAZAO_SOCIAL']}")
                if 'CNPJ_CONTRATANTE' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CNPJ_CONTRATANTE', f"{campos_contrato_final['CNPJ_CONTRATANTE']}")
                if 'CEP_EMPRESA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CEP_EMPRESA', f"{campos_contrato_final['CEP_EMPRESA']}")
                if 'ENDERECO_EMPRESA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('ENDERECO_EMPRESA', f"{campos_contrato_final['ENDERECO_EMPRESA']}")
                if 'NOME_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('NOME_AVALISTA', f"{campos_contrato_final['NOME_AVALISTA']}")
                if 'CPF_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CPF_AVALISTA', f"{campos_contrato_final['CPF_AVALISTA']}")
                if 'ENDERECO_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('ENDERECO_AVALISTA', f"{campos_contrato_final['ENDERECO_AVALISTA']}")
                if 'CEP_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CEP_AVALISTA', f"{campos_contrato_final['CEP_AVALISTA']}")
                if 'VALOR_MENSAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_MENSAL', f"{campos_contrato_final['VALOR_MENSAL']}")
                if 'VALOR_POR_EXTENSO' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_POR_EXTENSO', f"{campos_contrato_final['VALOR_POR_EXTENSO']}")
                if 'QTD_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_FUNC', f"{campos_contrato_final['QTD_FUNC']}")
                if 'VALOR_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_FUNC', f"{campos_contrato_final['VALOR_FUNC']}")
                if 'VALOR_POR_EXTENSO_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_POR_EXTENSO_FUNC', f"{campos_contrato_final['VALOR_POR_EXTENSO_FUNC']}")
                if 'REGIME' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('REGIME', f"{campos_contrato_final['REGIME']}")
                if 'QTD_FISCAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_FISCAL', f"{campos_contrato_final['QTD_FISCAL']}")
                if 'QTD_CONTABIL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_CONTABIL', f"{campos_contrato_final['QTD_CONTABIL']}")
                if 'VENCIMENTO' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VENCIMENTO', f"{campos_contrato_final['VENCIMENTO']}")
                if 'DATA_VIGENCIA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('DATA_VIGENCIA', f"{campos_contrato_final['DATA_VIGENCIA']}")
                if 'DATA_ASSINATURA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('DATA_ASSINATURA', f"{campos_contrato_final['DATA_ASSINATURA']}")
                if x < 18:
                    x += 1
                  
            con_domestica.save("Manipulação de contratos\ CONTRATO PREENCHIDO - DOMESTICA.docx")
def contrato_padrão():
    with col1:
        y = 0
        preencher = st.button(label='Preencher Campos')
        if preencher == True:
            for campo in campos_contrato:
                if campo == '':
                    st.error('Todos os campos devem estar preenchidos.')
                else:
                    campos_contrato_final[campos_contrato_strings[y]] = (campo)
                    y += 1
            x = 0
            for paragrafo in con_padrao.paragraphs:
                if '{{' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace("{{", '')
                if '}}' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace("}}", '')
                if 'RAZAO_SOCIAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('RAZAO_SOCIAL', f"{campos_contrato_final['RAZAO_SOCIAL']}")
                if 'CNPJ_CONTRATANTE' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CNPJ_CONTRATANTE', f"{campos_contrato_final['CNPJ_CONTRATANTE']}")
                if 'CEP_EMPRESA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CEP_EMPRESA', f"{campos_contrato_final['CEP_EMPRESA']}")
                if 'ENDERECO_EMPRESA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('ENDERECO_EMPRESA', f"{campos_contrato_final['ENDERECO_EMPRESA']}")
                if 'NOME_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('NOME_AVALISTA', f"{campos_contrato_final['NOME_AVALISTA']}")
                if 'CPF_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CPF_AVALISTA', f"{campos_contrato_final['CPF_AVALISTA']}")
                if 'ENDERECO_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('ENDERECO_AVALISTA', f"{campos_contrato_final['ENDERECO_AVALISTA']}")
                if 'CEP_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CEP_AVALISTA', f"{campos_contrato_final['CEP_AVALISTA']}")
                if 'VALOR_MENSAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_MENSAL', f"{campos_contrato_final['VALOR_MENSAL']}")
                if 'VALOR_POR_EXTENSO' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_POR_EXTENSO', f"{campos_contrato_final['VALOR_POR_EXTENSO']}")
                if 'QTD_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_FUNC', f"{campos_contrato_final['QTD_FUNC']}")
                if 'VALOR_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_FUNC', f"{campos_contrato_final['VALOR_FUNC']}")
                if 'VALOR_POR_EXTENSO_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_POR_EXTENSO_FUNC', f"{campos_contrato_final['VALOR_POR_EXTENSO_FUNC']}")
                if 'REGIME' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('REGIME', f"{campos_contrato_final['REGIME']}")
                if 'QTD_FISCAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_FISCAL', f"{campos_contrato_final['QTD_FISCAL']}")
                if 'QTD_CONTABIL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_CONTABIL', f"{campos_contrato_final['QTD_CONTABIL']}")
                if 'VENCIMENTO' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VENCIMENTO', f"{campos_contrato_final['VENCIMENTO']}")
                if 'DATA_VIGENCIA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('DATA_VIGENCIA', f"{campos_contrato_final['DATA_VIGENCIA']}")
                if 'DATA_ASSINATURA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('DATA_ASSINATURA', f"{campos_contrato_final['DATA_ASSINATURA']}")
                if x < 18:
                    x += 1

            con_padrao.save("Manipulação de contratos\ CONTRATO PREENCHIDO - PADRAO.docx")
def contrato_permuta():
    with col1:
        y = 0
        preencher = st.button(label='Preencher Campos')
        if preencher == True:
            for campo in campos_contrato:
                if campo == '':
                    st.error('Todos os campos devem estar preenchidos.')
                else:
                    campos_contrato_final[campos_contrato_strings[y]] = (campo)
                    y += 1
            x = 0
            for paragrafo in con_permuta.paragraphs:
                if '{{' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace("{{", '')
                if '}}' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace("}}", '')
                if 'RAZAO_SOCIAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('RAZAO_SOCIAL', f"{campos_contrato_final['RAZAO_SOCIAL']}")
                if 'CNPJ_CONTRATANTE' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CNPJ_CONTRATANTE', f"{campos_contrato_final['CNPJ_CONTRATANTE']}")
                if 'CEP_EMPRESA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CEP_EMPRESA', f"{campos_contrato_final['CEP_EMPRESA']}")
                if 'ENDERECO_EMPRESA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('ENDERECO_EMPRESA', f"{campos_contrato_final['ENDERECO_EMPRESA']}")
                if 'NOME_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('NOME_AVALISTA', f"{campos_contrato_final['NOME_AVALISTA']}")
                if 'CPF_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CPF_AVALISTA', f"{campos_contrato_final['CPF_AVALISTA']}")
                if 'ENDERECO_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('ENDERECO_AVALISTA', f"{campos_contrato_final['ENDERECO_AVALISTA']}")
                if 'CEP_AVALISTA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('CEP_AVALISTA', f"{campos_contrato_final['CEP_AVALISTA']}")
                if 'VALOR_MENSAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_MENSAL', f"{campos_contrato_final['VALOR_MENSAL']}")
                if 'VALOR_POR_EXTENSO' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_POR_EXTENSO', f"{campos_contrato_final['VALOR_POR_EXTENSO']}")
                if 'QTD_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_FUNC', f"{campos_contrato_final['QTD_FUNC']}")
                if 'VALOR_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_FUNC', f"{campos_contrato_final['VALOR_FUNC']}")
                if 'VALOR_POR_EXTENSO_FUNC' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VALOR_POR_EXTENSO_FUNC', f"{campos_contrato_final['VALOR_POR_EXTENSO_FUNC']}")
                if 'REGIME' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('REGIME', f"{campos_contrato_final['REGIME']}")
                if 'QTD_FISCAL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_FISCAL', f"{campos_contrato_final['QTD_FISCAL']}")
                if 'QTD_CONTABIL' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('QTD_CONTABIL', f"{campos_contrato_final['QTD_CONTABIL']}")
                if 'VENCIMENTO' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('VENCIMENTO', f"{campos_contrato_final['VENCIMENTO']}")
                if 'DATA_VIGENCIA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('DATA_VIGENCIA', f"{campos_contrato_final['DATA_VIGENCIA']}")
                if 'DATA_ASSINATURA' in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace('DATA_ASSINATURA', f"{campos_contrato_final['DATA_ASSINATURA']}")
                if x < 18:
                    x += 1

            con_permuta.save("Manipulação de contratos\ CONTRATO PREENCHIDO - PERMUTA.docx")

if ce == listaaquivos[0]:
    contrato_clube()
elif ce == listaaquivos[1]:
    contrato_domestica()
elif ce == listaaquivos[2]:
    contrato_padrão()
else:
    contrato_permuta()